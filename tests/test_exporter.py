from pathlib import Path

from docx import Document

from edu_question_generator.exporter import export_student_docx, export_teacher_docx
from edu_question_generator.generator import QuestionGenerator


def test_export_docx_files(tmp_path: Path):
    generator = QuestionGenerator(random_seed=3)
    worksheet = generator.generate("physics", "motion", "easy", 2)

    student_path = export_student_docx(worksheet, tmp_path / "student.docx")
    teacher_path = export_teacher_docx(worksheet, tmp_path / "teacher.docx")

    assert student_path.exists()
    assert teacher_path.exists()
    assert student_path.stat().st_size > 0
    assert teacher_path.stat().st_size > 0


def test_teacher_docx_contains_answers_but_student_docx_does_not(tmp_path: Path):
    generator = QuestionGenerator(random_seed=11)
    worksheet = generator.generate("physics", "motion", "easy", 1)

    student_path = export_student_docx(worksheet, tmp_path / "student.docx")
    teacher_path = export_teacher_docx(worksheet, tmp_path / "teacher.docx")

    student_text = "\n".join(p.text for p in Document(student_path).paragraphs)
    teacher_text = "\n".join(p.text for p in Document(teacher_path).paragraphs)

    assert worksheet.questions[0].prompt in student_text
    assert worksheet.questions[0].prompt in teacher_text
    assert worksheet.questions[0].answer not in student_text
    assert worksheet.questions[0].answer in teacher_text
