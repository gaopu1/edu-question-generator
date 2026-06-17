"""DOCX export helpers for student and teacher worksheet versions."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from .models import Worksheet


def export_student_docx(worksheet: Worksheet, path: str | Path) -> Path:
    """Export a worksheet without answers."""

    output_path = Path(path)
    document = _create_base_document(worksheet, "Student Worksheet")

    for index, question in enumerate(worksheet.questions, start=1):
        document.add_paragraph(f"{index}. {question.prompt}")
        document.add_paragraph("")
        document.add_paragraph("")

    document.save(output_path)
    return output_path


def export_teacher_docx(worksheet: Worksheet, path: str | Path) -> Path:
    """Export a teacher version with detailed answers."""

    output_path = Path(path)
    document = _create_base_document(worksheet, "Teacher Version")

    for index, question in enumerate(worksheet.questions, start=1):
        document.add_paragraph(f"{index}. {question.prompt}")
        answer_heading = document.add_paragraph()
        answer_heading.add_run("Answer: ").bold = True
        answer_heading.add_run(question.answer)

    document.save(output_path)
    return output_path


def _create_base_document(worksheet: Worksheet, subtitle: str) -> Document:
    document = Document()
    document.add_heading(worksheet.title, level=1)
    document.add_paragraph(subtitle)
    document.add_paragraph(f"Subject: {worksheet.subject.title()}")
    document.add_paragraph(f"Topic: {worksheet.topic.title()}")
    document.add_paragraph(f"Difficulty: {worksheet.difficulty.title()}")
    document.add_paragraph("")
    return document
